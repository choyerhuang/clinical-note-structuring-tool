from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ALLOWED_NOTE_EXTENSIONS = {".docx", ".pdf"}
SCANNED_PDF_WARNING = "No extractable text found in this PDF. It may be scanned."


def _normalize_text_lines(lines: list[str]) -> str:
    cleaned_lines = [line.strip() for line in lines if line and line.strip()]
    return "\n".join(cleaned_lines)


def _extract_docx_text(uploaded_file) -> str:
    uploaded_file.seek(0)
    document = Document(uploaded_file)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        lines.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)

    return _normalize_text_lines(lines)


def _extract_pdf_text(uploaded_file) -> str:
    uploaded_file.seek(0)
    pdf_bytes = uploaded_file.read()
    reader = PdfReader(BytesIO(pdf_bytes))
    lines: list[str] = []

    for page in reader.pages:
        lines.append(page.extract_text() or "")

    return _normalize_text_lines(lines)


def parse_uploaded_note_file(uploaded_file) -> dict:
    if uploaded_file is None:
        return {
            "success": False,
            "text": "",
            "warning": None,
            "error": "No file was provided.",
        }

    if getattr(uploaded_file, "size", 0) <= 0:
        return {
            "success": False,
            "text": "",
            "warning": None,
            "error": "The uploaded file is empty.",
        }

    extension = Path(uploaded_file.name or "").suffix.lower()
    if extension not in ALLOWED_NOTE_EXTENSIONS:
        return {
            "success": False,
            "text": "",
            "warning": None,
            "error": "Unsupported file type. Please upload a PDF or DOCX file.",
        }

    try:
        if extension == ".docx":
            text = _extract_docx_text(uploaded_file)
        else:
            text = _extract_pdf_text(uploaded_file)
    except Exception:
        return {
            "success": False,
            "text": "",
            "warning": None,
            "error": "Unable to parse file.",
        }

    if extension == ".pdf" and not text:
        return {
            "success": True,
            "text": "",
            "warning": SCANNED_PDF_WARNING,
            "error": None,
        }

    if not text:
        return {
            "success": False,
            "text": "",
            "warning": None,
            "error": "No extractable text was found in the uploaded file.",
        }

    return {
        "success": True,
        "text": text,
        "warning": None,
        "error": None,
    }
