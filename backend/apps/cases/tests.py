from io import BytesIO
from types import SimpleNamespace
from unittest.mock import patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import SimpleTestCase, TestCase
from docx import Document
from pypdf import PdfWriter
from rest_framework.test import APIClient

from apps.cases.services.compose import compose_revised_hpi
from apps.cases.services.criteria import (
    enrich_structured_output_with_source_evidence,
    match_mcg_criteria,
    reconcile_diabetes_disposition,
)
from apps.cases.services.extract import extract_structured_output
from apps.cases.services.file_parsing import parse_uploaded_note_file
from apps.cases.services.pipeline import run_generate_pipeline
from apps.cases.services.privacy import redact_phi
from apps.cases.services.validators import (
    build_general_missing_information,
    calculate_admission_support_confidence,
    build_generation_warning_groups,
    build_uncertainties,
    verify_revised_hpi,
)


class PrivacyRedactionTests(SimpleTestCase):
    def test_redact_phi_removes_labeled_phi_and_preserves_clinical_content(self):
        note = (
            "Patient Name: John Smith\n"
            "DOB: 01/23/1970\n"
            "MRN: 123456\n"
            "Patient ID: 998877\n"
            "Phone: (213) 555-1212\n"
            "Email: john@example.com\n"
            "SSN: 123-45-6789\n"
            "Address: 123 Main St, Los Angeles, CA\n"
            "53-year-old male presents with glucose 322 mg/dL.\n"
        )

        result = redact_phi(note)

        self.assertIn("Patient Name: [REDACTED_NAME]", result)
        self.assertIn("DOB: [REDACTED_DOB]", result)
        self.assertIn("MRN: [REDACTED_MRN]", result)
        self.assertIn("Patient ID: [REDACTED_PATIENT_ID]", result)
        self.assertIn("Phone: [REDACTED_PHONE]", result)
        self.assertIn("Email: [REDACTED_EMAIL]", result)
        self.assertIn("SSN: [REDACTED_SSN]", result)
        self.assertIn("Address: [REDACTED_ADDRESS]", result)
        self.assertIn("53-year-old male presents with glucose 322 mg/dL.", result)
        self.assertNotIn("John Smith", result)
        self.assertNotIn("john@example.com", result)
        self.assertNotIn("123-45-6789", result)


class ExtractStructuredOutputTests(SimpleTestCase):
    def test_extract_returns_valid_normalized_schema(self):
        fake_response = SimpleNamespace(
            output=[
                SimpleNamespace(
                    content=[
                        SimpleNamespace(
                            parsed={
                                "chief_complaint": "Chest pain",
                                "hpi_summary": "Acute chest pain with nausea.",
                                "key_findings": ["Chest pain", "Nausea"],
                                "suspected_conditions": ["ACS"],
                                "disposition_recommendation": "Observe",
                                "uncertainties": ["Troponin pending"],
                                "source_support": ["Chest pain for 2 hours"],
                            }
                        )
                    ]
                )
            ]
        )
        fake_provider = SimpleNamespace(
            generate_structured_json=lambda **kwargs: fake_response
        )

        with patch(
            "apps.cases.services.extract.get_llm_provider",
            return_value=(fake_provider, "fake-model"),
        ):
            result = extract_structured_output("Patient presents with chest pain.")

        self.assertEqual(result["chief_complaint_generated"], "Chest pain")
        self.assertEqual(result["disposition_generated"], "Observe")
        self.assertEqual(result["key_findings_generated"], ["Chest pain", "Nausea"])
        self.assertIn("source_support", result)


class VerifyRevisedHPITests(SimpleTestCase):
    def test_verify_returns_structured_json_shape(self):
        result = verify_revised_hpi(
            "The patient presents with chest pain and nausea for emergency evaluation.",
            {
                "chief_complaint_generated": "Chest pain",
                "hpi_summary_generated": "Chest pain with nausea.",
                "key_findings_generated": ["Chest pain", "Nausea"],
                "suspected_conditions_generated": [],
                "disposition_generated": "Unknown",
                "uncertainties_generated": ["Severity not documented."],
            },
            mcg_result={"matched_criteria": [], "support_level": "low", "supported": False, "missing_data": []},
        )

        self.assertEqual(
            {
                "is_pass",
                "factual_consistency",
                "requires_review",
                "unsupported_claims",
                "missing_key_facts",
                "disposition_consistency",
                "disposition_inconsistencies",
                "criteria_alignment_issues",
                "missing_required_data_for_confident_interpretation",
                "mcg_admission_check",
                "needs_regeneration",
                "revision_instructions",
            },
            set(result.keys()),
        )

    def test_verify_flags_unsupported_dka_claim(self):
        result = verify_revised_hpi(
            "The patient has diabetic ketoacidosis requiring admission.",
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Hyperglycemia with dehydration.",
                "key_findings_generated": ["Glucose 260 mg/dL", "Dehydration"],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Observe",
                "uncertainties_generated": ["pH unavailable"],
            },
            mcg_result={
                "matched_criteria": [],
                "support_level": "low",
                "supported": False,
                "missing_data": ["pH", "ketone level", "bicarbonate"],
            },
        )

        self.assertFalse(result["is_pass"])
        self.assertTrue(result["unsupported_claims"])
        self.assertTrue(result["needs_regeneration"])
        self.assertTrue(result["requires_review"])

    def test_verify_missing_data_only_returns_pass_with_review(self):
        result = verify_revised_hpi(
            "The patient presented with hyperglycemia and requires continued monitoring.",
            {
                "chief_complaint_generated": "Diabetes/Hyperglycemia",
                "hpi_summary_generated": "Hyperglycemia with dehydration and incomplete workup.",
                "key_findings_generated": ["Glucose 320 mg/dL", "Dehydration"],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            mcg_result={
                "applicable": True,
                "matched_criteria": [],
                "support_level": "low",
                "supported": False,
                "missing_data": ["ketone level", "pH"],
                "disposition_context": {
                    "requires_admit_with_monitoring": True,
                },
            },
        )

        self.assertTrue(result["is_pass"])
        self.assertTrue(result["requires_review"])
        self.assertEqual(result["factual_consistency"], "pass_with_warnings")
        self.assertFalse(result["needs_regeneration"])
        self.assertIn(
            "Chief complaint is present but generic.",
            result["missing_key_facts"],
        )


class SourceEvidenceEnrichmentTests(SimpleTestCase):
    def test_enrich_structured_output_adds_ketone_evidence_from_acetone_phrase(self):
        result = enrich_structured_output_with_source_evidence(
            {
                "chief_complaint_generated": "AMS, hyperglycemia",
                "hpi_summary_generated": "Hyperglycemia with acidosis concerns.",
                "key_findings_generated": ["Glucose 320 mg/dL"],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            source_text="Acetone Semiqt large A with glucose 320 mg/dL",
        )

        self.assertEqual(result["ketone_status"], "present")
        self.assertIn("Acetone Semiqt large", result["ketone_evidence"])
        self.assertIn("Large acetone/ketones present", result["key_findings_generated"])

    def test_match_mcg_does_not_mark_ketone_missing_when_acetone_is_present(self):
        result = match_mcg_criteria(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Hyperglycemia with dehydration.",
                "key_findings_generated": ["Glucose 320 mg/dL", "Large acetone/ketones present"],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            source_text="ACETONE LARGE A glucose 320 mg/dL bicarbonate 14 anion gap 18",
        )

        self.assertNotIn("ketone level", result["missing_data"])

    def test_verify_does_not_penalize_admit_when_diabetes_mcg_not_applicable(self):
        result = verify_revised_hpi(
            "The patient presents with fever, productive cough, and hypoxemia requiring admission for IV antibiotics and monitoring.",
            {
                "chief_complaint_generated": "Shortness of breath",
                "hpi_summary_generated": "Pneumonia with hypoxemia requiring inpatient care.",
                "key_findings_generated": [
                    "Fever",
                    "Productive cough",
                    "Hypoxemia",
                    "IV antibiotics started",
                ],
                "suspected_conditions_generated": ["pneumonia"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            mcg_result={
                "applicable": False,
                "matched_criteria": [],
                "support_level": "low",
                "supported": False,
                "missing_data": [],
            },
        )

        self.assertEqual(result["disposition_consistency"], "pass")
        self.assertFalse(result["disposition_inconsistencies"])


class AdmissionSupportConfidenceTests(SimpleTestCase):
    def test_calculate_admission_support_confidence_balances_support_and_missing_data(self):
        result = calculate_admission_support_confidence(
            {
                "is_pass": True,
                "factual_consistency": "pass",
                "disposition_consistency": "pass",
                "unsupported_claims": [],
                "disposition_inconsistencies": [],
                "missing_required_data_for_confident_interpretation": [
                    "ketone level",
                    "pH",
                    "bicarbonate",
                    "anion gap",
                    "serum osmolality",
                    "response to treatment",
                ],
            },
            {
                "applicable": True,
                "support_level": "moderate",
                "supported": True,
                "missing_data": [],
            },
            ["Missing data: pH"],
        )

        self.assertEqual(result["score"], 0.74)
        self.assertEqual(result["level"], "Medium")
        self.assertEqual(result["label"], "Admission Support Confidence")
        self.assertTrue(result["factors"])

    def test_confidence_penalizes_general_missing_information_without_collapsing_score(self):
        result = calculate_admission_support_confidence(
            {
                "is_pass": True,
                "factual_consistency": "pass",
                "disposition_consistency": "pass",
                "unsupported_claims": [],
                "disposition_inconsistencies": [],
                "missing_required_data_for_confident_interpretation": [],
            },
            {
                "applicable": False,
                "support_level": "low",
                "supported": False,
                "missing_data": [],
            },
            [
                "Missing data: Vital signs",
                "Missing data: Laboratory severity markers",
                "Missing data: Symptom duration",
                "Missing data: Response to prior treatment",
            ],
        )

        self.assertEqual(result["score"], 0.81)
        self.assertEqual(result["level"], "High")

    def test_confidence_applies_low_information_penalty(self):
        result = calculate_admission_support_confidence(
            {
                "is_pass": True,
                "factual_consistency": "pass",
                "disposition_consistency": "pass",
                "unsupported_claims": [],
                "disposition_inconsistencies": [],
                "missing_required_data_for_confident_interpretation": [],
            },
            {
                "applicable": False,
                "support_level": "low",
                "supported": False,
                "missing_data": [],
            },
            [],
            structured_output={
                "chief_complaint_generated": "Abdominal pain",
                "hpi_summary_generated": "",
                "key_findings_generated": ["Abdominal pain"],
                "suspected_conditions_generated": [],
                "disposition_generated": "Unknown",
                "uncertainties_generated": [],
            },
            source_text="Patient has abdominal pain.",
        )

        self.assertEqual(result["score"], 0.65)
        self.assertEqual(result["level"], "Medium")
        self.assertTrue(
            any(factor["label"] == "Low information penalty" for factor in result["factors"])
        )


class UncertaintyBuilderTests(SimpleTestCase):
    def test_build_general_missing_information_adds_non_condition_specific_gaps(self):
        result = build_general_missing_information(
            (
                "[ER NOTE]\n"
                "60-year-old female presents with abdominal pain and fever. CT suggests acute diverticulitis.\n\n"
                "[H&P NOTE]\n"
                "Patient is admitted for worsening symptoms and need for IV antibiotics."
            ),
            {
                "chief_complaint_generated": "Abdominal pain",
                "hpi_summary_generated": "Abdominal pain and fever with CT evidence of diverticulitis.",
                "key_findings_generated": [
                    "Abdominal pain",
                    "Fever",
                    "CT suggests acute diverticulitis",
                ],
                "suspected_conditions_generated": ["Diverticulitis"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
        )

        self.assertEqual(
            result,
            [
                "Vital signs not documented",
                "Key laboratory severity markers not available",
                "Duration of symptoms not specified",
                "Response to prior treatment not documented",
            ],
        )

    def test_build_uncertainties_merges_missingness_sources_without_duplicates(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Diabetes with dehydration.",
                "key_findings_generated": ["Glucose 320 mg/dL"],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": ["Serum osmolality not available"],
            },
            generation_warnings=["pH", "Response to treatment not yet documented"],
            mcg_result={"missing_data": ["serum osmolality", "post-treatment reassessment"]},
            verification={
                "missing_required_data_for_confident_interpretation": [
                    "pH",
                    "response to treatment",
                ]
            },
        )

        self.assertCountEqual(
            result,
            [
                "Serum osmolality not available",
                "pH not available to fully assess severity of acidosis",
                "Response to treatment not yet documented",
                "Post-treatment reassessment pending",
            ],
        )

    def test_build_uncertainties_leaves_clean_output_empty_when_no_missingness_exists(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Diabetes with ketonemia and metabolic acidosis.",
                "key_findings_generated": [
                    "Glucose 320 mg/dL",
                    "Anion gap 18",
                    "Bicarbonate 14",
                    "pH 7.21",
                    "Ketonemia",
                ],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            generation_warnings=[],
            mcg_result={"missing_data": []},
            verification={"missing_required_data_for_confident_interpretation": []},
        )

        self.assertEqual(result, [])

    def test_build_uncertainties_skips_false_glucose_missing_when_glucose_exists(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Diabetes with dehydration.",
                "key_findings_generated": ["Glucose 320 mg/dL", "Anion gap 18"],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            generation_warnings=["glucose"],
            mcg_result={"missing_data": ["glucose", "pH"]},
            verification={
                "missing_required_data_for_confident_interpretation": [
                    "glucose",
                    "pH",
                ]
            },
            source_text="glucose 320 mg/dL with dehydration and anion gap 18",
        )

        self.assertNotIn("Glucose value not available", result)
        self.assertIn(
            "pH not available to fully assess severity of acidosis",
            result,
        )

    def test_build_uncertainties_prefers_specific_ketone_phrase(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Diabetes evaluation.",
                "key_findings_generated": ["Glucose 280 mg/dL"],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": ["No ketone testing performed"],
            },
            mcg_result={"missing_data": ["ketone level"]},
            verification={},
            source_text="glucose 280 mg/dL",
        )

        self.assertIn("No ketone testing performed", result)
        self.assertNotIn("Ketone status not available", result)

    def test_build_uncertainties_splits_combined_ph_and_bicarbonate_phrase(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Diabetes with incomplete lab assessment.",
                "key_findings_generated": ["Glucose 260 mg/dL"],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Observe",
                "uncertainties_generated": ["Bicarbonate and pH not available"],
            },
            mcg_result={"missing_data": ["pH", "bicarbonate"]},
            verification={},
            source_text="glucose 260 mg/dL",
        )

        self.assertCountEqual(
            result,
            [
                "Bicarbonate level not available",
                "pH not available to fully assess severity of acidosis",
            ],
        )

    def test_build_uncertainties_filters_out_verifier_and_qa_text(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Shortness of breath",
                "hpi_summary_generated": "Pneumonia requiring admission.",
                "key_findings_generated": ["Hypoxemia", "IV antibiotics started"],
                "suspected_conditions_generated": ["pneumonia"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [
                    "Revised HPI may need regeneration",
                    "Chief complaint is not clearly stated.",
                    "Serum osmolality not available",
                ],
            },
            generation_warnings=[
                "Admission disposition should more clearly explain the need for ongoing monitoring or incomplete evaluation when condition-specific criteria support is limited."
            ],
            mcg_result={"missing_data": []},
            verification={
                "missing_required_data_for_confident_interpretation": [],
            },
            source_text="pneumonia with hypoxemia",
        )

        self.assertEqual(result, ["Serum osmolality not available"])

    def test_build_uncertainties_final_consistency_removes_present_findings(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Diabetes with ketones and anion gap elevation.",
                "key_findings_generated": [
                    "Glucose 320 mg/dL",
                    "Ketones present",
                    "Anion gap 18",
                ],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [
                    "Ketone status not available",
                    "Anion gap not available",
                    "Glucose value not available",
                    "pH not available to fully assess severity of acidosis",
                ],
            },
            mcg_result={
                "missing_data": ["ketone level", "anion gap", "glucose", "pH"],
            },
            verification={
                "missing_required_data_for_confident_interpretation": [
                    "ketone level",
                    "anion gap",
                    "glucose",
                    "pH",
                ]
            },
            source_text="glucose 320 mg/dL, ketones present, anion gap 18",
        )

        self.assertEqual(
            result,
            ["pH not available to fully assess severity of acidosis"],
        )

    def test_build_uncertainties_includes_general_missing_information_for_diverticulitis(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Abdominal pain",
                "hpi_summary_generated": "Abdominal pain and fever with CT evidence of diverticulitis.",
                "key_findings_generated": [
                    "Abdominal pain",
                    "Fever",
                    "CT suggests acute diverticulitis",
                ],
                "suspected_conditions_generated": ["Diverticulitis"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            mcg_result={"missing_data": []},
            verification={"missing_required_data_for_confident_interpretation": []},
            source_text=(
                "[ER NOTE]\n"
                "60-year-old female presents with abdominal pain and fever. CT suggests acute diverticulitis.\n\n"
                "[H&P NOTE]\n"
                "Patient is admitted for worsening symptoms and need for IV antibiotics."
            ),
        )

        self.assertEqual(
            result,
            [
                "Vital signs not documented",
                "Key laboratory severity markers not available",
                "Duration of symptoms not specified",
                "Response to prior treatment not documented",
            ],
        )

    def test_build_uncertainties_merges_general_missing_items_from_generation_warnings(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Abdominal pain",
                "hpi_summary_generated": "Acute diverticulitis admitted for IV antibiotics.",
                "key_findings_generated": ["CT suggests acute diverticulitis"],
                "suspected_conditions_generated": ["Diverticulitis"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            generation_warnings=[
                "Missing data: Vital signs",
                "Missing data: Laboratory severity markers",
                "Missing data: Symptom duration",
                "Missing data: Response to prior treatment",
            ],
            mcg_result={"missing_data": []},
            verification={"missing_required_data_for_confident_interpretation": []},
            source_text=(
                "[ER NOTE]\n"
                "60-year-old female presents with abdominal pain and fever. CT suggests acute diverticulitis.\n\n"
                "[H&P NOTE]\n"
                "Patient is admitted for worsening symptoms and need for IV antibiotics."
            ),
        )

        self.assertEqual(
            result,
            [
                "Vital signs not documented",
                "Key laboratory severity markers not available",
                "Duration of symptoms not specified",
                "Response to prior treatment not documented",
            ],
        )

    def test_build_uncertainties_treats_qualitative_glucose_and_positive_ketones_as_present(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Markedly elevated glucose with positive ketones.",
                "key_findings_generated": [
                    "Markedly elevated glucose",
                    "Positive ketones",
                ],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [
                    "Glucose value not available",
                    "Ketone status not available",
                    "pH not available to fully assess severity of acidosis",
                ],
            },
            mcg_result={"missing_data": ["glucose", "ketone level", "pH"]},
            verification={
                "missing_required_data_for_confident_interpretation": [
                    "glucose",
                    "ketone level",
                    "pH",
                ]
            },
            source_text="markedly elevated glucose with positive ketones",
        )

        self.assertEqual(
            result,
            ["pH not available to fully assess severity of acidosis"],
        )

    def test_build_uncertainties_treats_qualitative_elevated_anion_gap_as_present(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Diabetes with elevated anion gap.",
                "key_findings_generated": [
                    "Elevated anion gap",
                    "Positive ketones",
                ],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [
                    "Anion gap not available",
                    "Ketone status not available",
                    "pH not available to fully assess severity of acidosis",
                ],
            },
            mcg_result={"missing_data": ["anion gap", "ketone level", "pH"]},
            verification={
                "missing_required_data_for_confident_interpretation": [
                    "anion gap",
                    "ketone level",
                    "pH",
                ]
            },
            source_text="elevated anion gap with positive ketones",
        )

        self.assertEqual(
            result,
            ["pH not available to fully assess severity of acidosis"],
        )

    def test_build_uncertainties_treats_documented_ph_as_present(self):
        result = build_uncertainties(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Diabetes with metabolic acidosis.",
                "key_findings_generated": [
                    "Glucose 320 mg/dL",
                    "pH 7.24",
                ],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [
                    "pH not available to fully assess severity of acidosis",
                    "Bicarbonate level not available",
                ],
            },
            mcg_result={"missing_data": ["pH", "bicarbonate"]},
            verification={
                "missing_required_data_for_confident_interpretation": ["pH", "bicarbonate"]
            },
            source_text="glucose 320 mg/dL and pH 7.24",
        )

        self.assertEqual(result, ["Bicarbonate level not available"])

    def test_build_generation_warning_groups_returns_display_friendly_labels(self):
        warning_groups = build_generation_warning_groups(
            ["Structured output is sparse and may need clinician review."],
            {
                "factual_consistency": "fail",
                "disposition_consistency": "fail",
                "needs_regeneration": True,
                "unsupported_claims": [],
                "missing_key_facts": ["Chief complaint is not clearly stated."],
                "disposition_inconsistencies": [],
                "criteria_alignment_issues": [],
                "missing_required_data_for_confident_interpretation": [
                    "ketone level",
                    "pH",
                    "response to treatment",
                ],
            },
            mcg_result={"missing_data": ["anion gap", "ketone level"]},
        )

        self.assertEqual(
            warning_groups["missing_data"],
            ["Anion gap", "Ketone level", "pH", "Response to treatment documentation"],
        )
        self.assertIn(
            "Chief complaint may not be clearly defined.",
            warning_groups["potential_issues"],
        )
        self.assertIn(
            "Revised HPI should be reviewed for consistency.",
            warning_groups["potential_issues"],
        )


class CriteriaMatchingTests(SimpleTestCase):
    def test_dka_case_matches_high_support(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Diabetes with nausea and dehydration.",
            "key_findings_generated": [
                "Glucose 320 mg/dL",
                "Anion gap 18",
                "Bicarbonate 14",
                "Ketones present",
                "Poor oral intake",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": ["pH not available"],
        }
        result = match_mcg_criteria(structured, source_text="glucose 320 anion gap 18 bicarbonate 14 ketones present poor oral intake")

        self.assertEqual(result["support_level"], "high")
        self.assertTrue(
            any(item["id"] == "diabetic_ketoacidosis" for item in result["matched_criteria"])
        )
        self.assertNotIn("glucose", result["missing_data"])

    def test_incomplete_hyperglycemia_does_not_overcall_dka(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Diabetes with dehydration.",
            "key_findings_generated": ["Glucose 260 mg/dL", "Dehydration"],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Unknown",
            "uncertainties_generated": [],
        }
        result = match_mcg_criteria(structured, source_text="glucose 260 dehydration")

        self.assertFalse(
            any(item["id"] == "diabetic_ketoacidosis" for item in result["matched_criteria"])
        )
        self.assertIn("pH", result["missing_data"])
        self.assertIn("ketone level", result["missing_data"])
        self.assertNotIn("glucose", result["missing_data"])

    def test_missing_data_final_consistency_removes_present_findings(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Diabetes with ketones and anion gap elevation.",
            "key_findings_generated": [
                "Glucose 320 mg/dL",
                "Ketones present",
                "Anion gap 18",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": [],
        }
        result = match_mcg_criteria(
            structured,
            source_text="glucose 320 mg/dL, ketones present, anion gap 18",
        )

        self.assertNotIn("glucose", result["missing_data"])
        self.assertNotIn("ketone level", result["missing_data"])
        self.assertNotIn("anion gap", result["missing_data"])

    def test_failed_observation_matches_treatment_failure(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Persistent dehydration and glucose not stabilized after treatment.",
            "key_findings_generated": ["Persistent dehydration", "Glucose not stabilized after treatment"],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": [],
        }
        result = match_mcg_criteria(
            structured,
            source_text="persistent dehydration and glucose not stabilized after treatment",
        )

        self.assertTrue(
            any(item["id"] == "failed_observation_or_outpatient" for item in result["matched_criteria"])
        )
        self.assertIn(result["support_level"], {"moderate", "high"})

    def test_failed_observation_recognizes_persistent_hyperglycemia_despite_insulin(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Persistent hyperglycemia despite insulin with dehydration despite fluids.",
            "key_findings_generated": [
                "Persistent hyperglycemia despite insulin",
                "Symptoms not improved significantly",
                "IV fluids started",
                "Dehydration persists",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": [],
        }
        result = match_mcg_criteria(
            structured,
            source_text=(
                "persistent hyperglycemia despite insulin, symptoms not improved significantly, "
                "IV fluids started, dehydration persists"
            ),
        )

        self.assertTrue(
            any(item["id"] == "failed_observation_or_outpatient" for item in result["matched_criteria"])
        )
        self.assertIn(result["support_level"], {"moderate", "high"})
        self.assertNotIn(
            "limited because key metabolic data are unavailable",
            result["criteria_summary"].lower(),
        )
        self.assertIn("persistent hyperglycemia", result["criteria_summary"].lower())
        self.assertNotIn("were not clearly met", result["criteria_summary"].lower())

    def test_failed_observation_with_unchanged_symptoms_and_monitoring_supports_admission(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Persistent hyperglycemia despite insulin with dehydration and no significant improvement after ED treatment.",
            "key_findings_generated": [
                "Persistent hyperglycemia despite insulin",
                "Dehydration",
                "No significant improvement after ED treatment",
                "Need for continued monitoring",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": [],
        }
        result = match_mcg_criteria(
            structured,
            source_text=(
                "persistent hyperglycemia despite insulin, dehydration, no significant improvement "
                "after ED treatment, need for continued monitoring"
            ),
        )

        self.assertTrue(result["supported"])
        self.assertIn(result["support_level"], {"moderate", "high"})
        self.assertIn("monitoring", result["criteria_summary"].lower())

    def test_underlying_inpatient_condition_matches(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Diabetes with sepsis requiring IV treatment.",
            "key_findings_generated": ["Sepsis", "IV antibiotics"],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": [],
        }
        result = match_mcg_criteria(
            structured,
            source_text="diabetes plus sepsis requiring IV treatment",
        )

        self.assertTrue(
            any(
                item["id"] == "underlying_condition_requires_inpatient"
                for item in result["matched_criteria"]
            )
        )

    def test_non_diabetes_note_does_not_match_diabetes_criteria(self):
        structured = {
            "chief_complaint_generated": "Chest pain",
            "hpi_summary_generated": "Chest pain evaluation.",
            "key_findings_generated": ["Chest pain"],
            "suspected_conditions_generated": ["acute coronary syndrome"],
            "disposition_generated": "Observe",
            "uncertainties_generated": [],
        }
        result = match_mcg_criteria(structured, source_text="chest pain without diabetes history")

        self.assertFalse(result["matched_criteria"])
        self.assertEqual(result["support_level"], "low")
        self.assertFalse(result["applicable"])
        self.assertEqual(result["criteria_summary"], "")


class DiabetesDispositionConsistencyTests(SimpleTestCase):
    def test_low_support_incomplete_diabetes_case_defaults_to_observe_without_inpatient_need(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Diabetes with dehydration and incomplete DKA evaluation.",
            "key_findings_generated": ["Glucose 260 mg/dL", "Dehydration"],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": ["Ketone status not available"],
        }
        mcg_result = match_mcg_criteria(structured, source_text="glucose 260 dehydration")

        reconciled_structured, reconciled_mcg = reconcile_diabetes_disposition(
            structured,
            mcg_result,
            source_text="glucose 260 dehydration",
        )

        self.assertEqual(mcg_result["support_level"], "low")
        self.assertEqual(reconciled_structured["disposition_generated"], "Observe")
        self.assertTrue(reconciled_mcg["disposition_adjusted_to_observe"])

    def test_low_support_incomplete_diabetes_case_can_remain_admit_with_monitoring_need(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Diabetes with dehydration and incomplete workup requiring continued monitoring.",
            "key_findings_generated": [
                "Glucose 260 mg/dL",
                "Dehydration",
                "Need for continued monitoring",
                "Workup remains incomplete",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": ["Ketone status not available"],
        }
        mcg_result = match_mcg_criteria(
            structured,
            source_text="glucose 260 dehydration workup remains incomplete and need for continued monitoring",
        )

        reconciled_structured, reconciled_mcg = reconcile_diabetes_disposition(
            structured,
            mcg_result,
            source_text="glucose 260 dehydration workup remains incomplete and need for continued monitoring",
        )

        self.assertEqual(reconciled_structured["disposition_generated"], "Admit")
        self.assertTrue(
            mcg_result["support_level"] in {"moderate", "high"}
            or reconciled_mcg["disposition_context"]["requires_admit_with_monitoring"]
        )

    def test_treatment_failure_diabetes_case_promotes_unknown_disposition_to_admit(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Persistent symptoms and no improvement after prior treatment.",
            "key_findings_generated": [
                "Persistent hyperglycemia despite insulin",
                "Symptoms not improved significantly",
                "Need for continued monitoring",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Unknown",
            "uncertainties_generated": ["pH not available to fully assess severity of acidosis"],
        }
        mcg_result = match_mcg_criteria(
            structured,
            source_text=(
                "persistent hyperglycemia despite insulin, symptoms not improved significantly, "
                "need for continued monitoring"
            ),
        )

        reconciled_structured, reconciled_mcg = reconcile_diabetes_disposition(
            structured,
            mcg_result,
            source_text=(
                "persistent hyperglycemia despite insulin, symptoms not improved significantly, "
                "need for continued monitoring"
            ),
        )

        self.assertIn(mcg_result["support_level"], {"moderate", "high"})
        self.assertEqual(reconciled_structured["disposition_generated"], "Admit")
        self.assertTrue(reconciled_mcg["disposition_inferred_from_treatment_failure"])

    def test_treatment_failure_diabetes_case_promotes_observe_to_admit(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Persistent symptoms with no improvement after prior insulin treatment.",
            "key_findings_generated": [
                "Persistent symptoms",
                "No improvement after prior treatment",
                "Insulin given earlier",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Observe",
            "uncertainties_generated": ["pH not available to fully assess severity of acidosis"],
        }
        mcg_result = match_mcg_criteria(
            structured,
            source_text=(
                "persistent symptoms, no improvement after prior treatment, insulin given earlier"
            ),
        )

        reconciled_structured, reconciled_mcg = reconcile_diabetes_disposition(
            structured,
            mcg_result,
            source_text=(
                "persistent symptoms, no improvement after prior treatment, insulin given earlier"
            ),
        )

        self.assertIn(mcg_result["support_level"], {"moderate", "high"})
        self.assertNotIn("criteria not clearly met", mcg_result["criteria_summary"].lower())
        self.assertEqual(reconciled_structured["disposition_generated"], "Admit")
        self.assertTrue(reconciled_mcg["disposition_inferred_from_treatment_failure"])
        self.assertIn(
            "failure of lower-level care",
            reconciled_mcg["criteria_summary"].lower(),
        )

    def test_treatment_failure_with_ongoing_symptoms_and_poor_intake_supports_admission(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Persistent symptoms, poor intake, and no improvement after insulin.",
            "key_findings_generated": [
                "Persistent symptoms",
                "No improvement after prior treatment",
                "Insulin given earlier",
                "Poor oral intake",
                "Dehydration",
            ],
            "suspected_conditions_generated": ["diabetes-related complications"],
            "disposition_generated": "Observe",
            "uncertainties_generated": [
                "Ketone status not available",
                "pH not available to fully assess severity of acidosis",
            ],
        }
        mcg_result = match_mcg_criteria(
            structured,
            source_text=(
                "persistent symptoms, no improvement after prior treatment, insulin given earlier, "
                "poor oral intake, dehydration"
            ),
        )

        reconciled_structured, reconciled_mcg = reconcile_diabetes_disposition(
            structured,
            mcg_result,
            source_text=(
                "persistent symptoms, no improvement after prior treatment, insulin given earlier, "
                "poor oral intake, dehydration"
            ),
        )

        self.assertIn(reconciled_mcg["support_level"], {"moderate", "high"})
        self.assertEqual(reconciled_structured["disposition_generated"], "Admit")
        self.assertIn("continued inpatient monitoring", reconciled_mcg["criteria_summary"].lower())
        self.assertNotIn(
            "limited because key metabolic data are unavailable",
            reconciled_mcg["criteria_summary"].lower(),
        )

    def test_mild_stable_diabetes_case_avoids_admission(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Mild hyperglycemia with otherwise stable presentation.",
            "key_findings_generated": [
                "Mild hyperglycemia",
                "Stable vital signs",
                "No dehydration",
                "Normal mental status",
                "No concerning symptoms",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": [],
        }
        mcg_result = match_mcg_criteria(
            structured,
            source_text=(
                "mild hyperglycemia, stable vital signs, no dehydration, normal mental status, "
                "no concerning symptoms"
            ),
        )

        reconciled_structured, reconciled_mcg = reconcile_diabetes_disposition(
            structured,
            mcg_result,
            source_text=(
                "mild hyperglycemia, stable vital signs, no dehydration, normal mental status, "
                "no concerning symptoms"
            ),
        )

        self.assertEqual(reconciled_structured["disposition_generated"], "Discharge")
        self.assertTrue(reconciled_mcg["mild_stable_outpatient_candidate"])
        self.assertEqual(reconciled_mcg["support_level"], "low")
        self.assertFalse(reconciled_mcg["supported"])
        self.assertIn("inpatient-level care is not required", reconciled_mcg["criteria_summary"].lower())
        self.assertTrue(
            reconciled_mcg["disposition_context"]["inpatient_level_care_not_required"]
        )

    def test_unknown_disposition_is_not_promoted_without_explicit_treatment_failure(self):
        structured = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Hyperglycemia without clear improvement data.",
            "key_findings_generated": [
                "Glucose remains elevated",
                "Dehydration",
            ],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Unknown",
            "uncertainties_generated": [],
        }
        mcg_result = match_mcg_criteria(
            structured,
            source_text="glucose remains elevated and dehydration",
        )

        reconciled_structured, reconciled_mcg = reconcile_diabetes_disposition(
            structured,
            mcg_result,
            source_text="glucose remains elevated and dehydration",
        )

        self.assertEqual(reconciled_structured["disposition_generated"], "Unknown")
        self.assertNotIn("disposition_inferred_from_treatment_failure", reconciled_mcg)


class ComposeRevisedHPITests(SimpleTestCase):
    @patch("apps.cases.services.compose.get_llm_provider")
    def test_compose_omits_admission_support_context_when_mcg_not_applicable(self, mock_get_llm_provider):
        captured = {}

        def fake_generate_text(**kwargs):
            captured["user_prompt"] = kwargs["user_prompt"]
            return SimpleNamespace(output_text="Pneumonia admission narrative.")

        fake_provider = SimpleNamespace(generate_text=fake_generate_text)
        mock_get_llm_provider.return_value = (fake_provider, "fake-model")

        result = compose_revised_hpi(
            {
                "chief_complaint_generated": "Shortness of breath",
                "hpi_summary_generated": "Pneumonia with hypoxemia.",
                "key_findings_generated": ["Fever", "Hypoxemia", "IV antibiotics started"],
                "suspected_conditions_generated": ["pneumonia"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [],
            },
            mcg_result={
                "applicable": False,
                "matched_criteria": [],
                "support_level": "low",
                "criteria_summary": "",
                "supported": False,
                "missing_data": [],
            },
        )

        self.assertEqual(result, "Pneumonia admission narrative.")
        self.assertNotIn("admission_support_context", captured["user_prompt"])
        self.assertIn('"care_plan_context"', captured["user_prompt"])
        self.assertIn("continued IV antibiotic therapy", captured["user_prompt"])

    @patch("apps.cases.services.compose.get_llm_provider")
    def test_compose_marks_treatment_failure_as_independent_admission_support_branch(self, mock_get_llm_provider):
        captured = {}

        def fake_generate_text(**kwargs):
            captured["user_prompt"] = kwargs["user_prompt"]
            return SimpleNamespace(output_text="Treatment-failure diabetes admission narrative.")

        fake_provider = SimpleNamespace(generate_text=fake_generate_text)
        mock_get_llm_provider.return_value = (fake_provider, "fake-model")

        result = compose_revised_hpi(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Persistent hyperglycemia despite insulin with dehydration and poor response to ED treatment.",
                "key_findings_generated": [
                    "Persistent hyperglycemia despite insulin",
                    "Dehydration",
                    "No significant improvement after ED treatment",
                    "Need for continued monitoring",
                ],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [
                    "Ketone status not available",
                    "pH not available to fully assess severity of acidosis",
                ],
            },
            mcg_result={
                "applicable": True,
                "matched_criteria": [
                    {
                        "id": "failed_observation_or_outpatient",
                        "matched_signals": [
                            "glucose not stabilized after treatment",
                            "failed initial management",
                            "continued monitoring or treatment required",
                        ],
                    },
                    {
                        "id": "severe_hyperglycemia",
                        "matched_signals": [
                            "persistent hyperglycemia despite insulin",
                            "dehydration",
                        ],
                    },
                ],
                "support_level": "moderate",
                "criteria_summary": (
                    "Admission support is present based on persistent hyperglycemia, "
                    "treatment failure, dehydration, or the need for continued monitoring and management."
                ),
                "supported": True,
                "missing_data": ["ketone level", "pH", "bicarbonate", "anion gap"],
            },
        )

        self.assertEqual(result, "Treatment-failure diabetes admission narrative.")
        self.assertIn('"support_branch": "treatment_failure"', captured["user_prompt"])
        self.assertIn("failure of lower-level care", captured["user_prompt"])
        self.assertIn("continued glucose control", captured["user_prompt"])
        self.assertIn("ongoing hydration support", captured["user_prompt"])
        self.assertNotIn('"missing_data": ["ketone level", "pH", "bicarbonate", "anion gap"]', captured["user_prompt"])

    @patch("apps.cases.services.compose.get_llm_provider")
    def test_compose_includes_admit_monitoring_context_for_incomplete_low_support_diabetes_case(self, mock_get_llm_provider):
        captured = {}

        def fake_generate_text(**kwargs):
            captured["user_prompt"] = kwargs["user_prompt"]
            return SimpleNamespace(output_text="Incomplete diabetes admit narrative.")

        fake_provider = SimpleNamespace(generate_text=fake_generate_text)
        mock_get_llm_provider.return_value = (fake_provider, "fake-model")

        result = compose_revised_hpi(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Hyperglycemia with dehydration and incomplete workup.",
                "key_findings_generated": [
                    "Glucose 260 mg/dL",
                    "Dehydration",
                    "Need for continued monitoring",
                ],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Admit",
                "uncertainties_generated": [
                    "Ketone status not available",
                    "pH not available to fully assess severity of acidosis",
                ],
            },
            mcg_result={
                "applicable": True,
                "matched_criteria": [],
                "support_level": "low",
                "criteria_summary": "",
                "supported": False,
                "missing_data": ["ketone level", "pH", "bicarbonate"],
                "disposition_context": {
                    "requires_admit_with_monitoring": True,
                    "guidance": (
                        "Admission remains appropriate because the workup is incomplete or still "
                        "concerning, and the patient requires ongoing monitoring, reassessment, "
                        "or treatment that cannot yet be stepped down."
                    ),
                },
            },
        )

        self.assertEqual(result, "Incomplete diabetes admit narrative.")
        self.assertIn('"requires_admit_with_monitoring": true', captured["user_prompt"])
        self.assertIn("workup is incomplete or still concerning", captured["user_prompt"])

    @patch("apps.cases.services.compose.get_llm_provider")
    def test_compose_includes_outpatient_context_for_mild_stable_diabetes_case(self, mock_get_llm_provider):
        captured = {}

        def fake_generate_text(**kwargs):
            captured["user_prompt"] = kwargs["user_prompt"]
            return SimpleNamespace(output_text="Mild stable diabetes outpatient narrative.")

        fake_provider = SimpleNamespace(generate_text=fake_generate_text)
        mock_get_llm_provider.return_value = (fake_provider, "fake-model")

        result = compose_revised_hpi(
            {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Mild hyperglycemia with otherwise stable presentation.",
                "key_findings_generated": [
                    "Mild hyperglycemia",
                    "Stable vital signs",
                    "No dehydration",
                    "Normal mental status",
                ],
                "suspected_conditions_generated": ["diabetes"],
                "disposition_generated": "Discharge",
                "uncertainties_generated": [],
            },
            mcg_result={
                "applicable": True,
                "matched_criteria": [],
                "support_level": "low",
                "criteria_summary": "",
                "supported": False,
                "missing_data": [],
                "disposition_context": {
                    "inpatient_level_care_not_required": True,
                    "guidance": (
                        "The available evidence suggests mild hyperglycemia without instability, "
                        "metabolic derangement, dehydration, or concerning symptoms, so inpatient-level "
                        "care is not required and discharge with outpatient follow-up is appropriate."
                    ),
                },
            },
        )

        self.assertEqual(result, "Mild stable diabetes outpatient narrative.")
        self.assertIn('"inpatient_level_care_not_required": true', captured["user_prompt"])
        self.assertIn("discharge with outpatient follow-up is appropriate", captured["user_prompt"])
        self.assertIn('"support_branch": "outpatient_management"', captured["user_prompt"])
        self.assertIn("inpatient-level care is not indicated", captured["user_prompt"].lower())
        self.assertIn("do not describe instability, treatment failure, need for inpatient monitoring, or admission support", captured["user_prompt"].lower())


class GeneratePipelineTests(SimpleTestCase):
    @patch("apps.cases.services.pipeline.verify_revised_hpi")
    @patch("apps.cases.services.pipeline.compose_revised_hpi")
    @patch("apps.cases.services.pipeline.match_mcg_criteria")
    @patch("apps.cases.services.pipeline.extract_structured_output")
    def test_pipeline_uses_redacted_note_for_llm_and_original_note_for_local_logic(
        self,
        mock_extract,
        mock_match_mcg,
        mock_compose,
        mock_verify,
    ):
        original_note = (
            "Patient Name: John Smith\n"
            "DOB: 01/23/1970\n"
            "MRN: 123456\n"
            "Phone: (213) 555-1212\n"
            "53-year-old male presents with glucose 322 mg/dL.\n"
        )
        mock_extract.return_value = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Hyperglycemia requiring reassessment.",
            "key_findings_generated": ["Glucose 322 mg/dL"],
            "suspected_conditions_generated": ["Hyperglycemia"],
            "disposition_generated": "Observe",
            "uncertainties_generated": [],
        }
        mock_match_mcg.return_value = {
            "condition": "diabetes",
            "matched_criteria": [],
            "support_level": "low",
            "missing_data": [],
            "criteria_summary": "",
            "supported": False,
        }
        mock_compose.return_value = "Revised HPI."
        mock_verify.return_value = {
            "factual_consistency": "pass",
            "unsupported_claims": [],
            "missing_key_facts": [],
            "disposition_consistency": "pass",
            "disposition_inconsistencies": [],
            "criteria_alignment_issues": [],
            "missing_required_data_for_confident_interpretation": [],
            "mcg_admission_check": {
                "matched_criteria": [],
                "support_level": "low",
                "supported": False,
            },
            "needs_regeneration": False,
            "revision_instructions": [],
            "is_pass": True,
        }

        result = run_generate_pipeline(original_note)

        extracted_note = mock_extract.call_args.args[0]
        self.assertIn("[REDACTED_NAME]", extracted_note)
        self.assertIn("[REDACTED_DOB]", extracted_note)
        self.assertIn("[REDACTED_MRN]", extracted_note)
        self.assertIn("[REDACTED_PHONE]", extracted_note)
        self.assertIn("53-year-old male presents with glucose 322 mg/dL.", extracted_note)
        self.assertEqual(mock_match_mcg.call_args.kwargs["source_text"], original_note)
        self.assertTrue(result["privacy"]["phi_redaction_applied"])
        self.assertIn("NAME", result["privacy"]["redaction_types"])

    @patch("apps.cases.services.pipeline.verify_revised_hpi")
    @patch("apps.cases.services.pipeline.compose_revised_hpi")
    @patch("apps.cases.services.pipeline.match_mcg_criteria")
    @patch("apps.cases.services.pipeline.extract_structured_output")
    def test_pipeline_retries_once_when_verification_requests_regeneration(
        self,
        mock_extract,
        mock_match_mcg,
        mock_compose,
        mock_verify,
    ):
        mock_extract.return_value = {
            "chief_complaint_generated": "Chest pain",
            "hpi_summary_generated": "Chest pain with nausea.",
            "key_findings_generated": ["Chest pain"],
            "suspected_conditions_generated": [],
            "disposition_generated": "Unknown",
            "uncertainties_generated": ["Severity unclear."],
        }
        mock_match_mcg.return_value = {
            "condition": "diabetes",
            "matched_criteria": [],
            "support_level": "low",
            "missing_data": [],
            "criteria_summary": "",
            "supported": False,
        }
        mock_compose.side_effect = [
            "First revised HPI.",
            "Second revised HPI.",
        ]
        mock_verify.side_effect = [
            {
                "factual_consistency": "fail",
                "unsupported_claims": ["Unsupported statement."],
                "missing_key_facts": [],
                "disposition_consistency": "pass",
                "needs_regeneration": True,
                "revision_instructions": ["Remove unsupported statement."],
            },
            {
                "factual_consistency": "pass",
                "unsupported_claims": [],
                "missing_key_facts": [],
                "disposition_consistency": "pass",
                "needs_regeneration": False,
                "revision_instructions": [],
            },
        ]

        result = run_generate_pipeline("Patient presents with chest pain.")

        self.assertEqual(mock_compose.call_count, 2)
        self.assertEqual(result["revised_hpi"], "Second revised HPI.")
        self.assertEqual(result["verification"]["needs_regeneration"], False)
        self.assertIn("mcg_result", result)
        self.assertIn("warning_groups", result)

    @patch("apps.cases.services.pipeline.verify_revised_hpi")
    @patch("apps.cases.services.pipeline.compose_revised_hpi")
    @patch("apps.cases.services.pipeline.match_mcg_criteria")
    @patch("apps.cases.services.pipeline.extract_structured_output")
    def test_pipeline_populates_uncertainties_from_known_missingness_signals(
        self,
        mock_extract,
        mock_match_mcg,
        mock_compose,
        mock_verify,
    ):
        mock_extract.return_value = {
            "chief_complaint_generated": "Hyperglycemia",
            "hpi_summary_generated": "Diabetes with dehydration.",
            "key_findings_generated": ["Glucose 320 mg/dL", "Anion gap 18"],
            "suspected_conditions_generated": ["diabetes"],
            "disposition_generated": "Admit",
            "uncertainties_generated": [],
        }
        mock_match_mcg.return_value = {
            "condition": "diabetes",
            "matched_criteria": [{"id": "diabetic_ketoacidosis", "confidence": "high"}],
            "support_level": "high",
            "missing_data": ["pH", "serum osmolality", "response to treatment"],
            "criteria_summary": "Admission is supported by DKA physiology.",
            "supported": True,
        }
        mock_compose.return_value = "Revised HPI."
        mock_verify.return_value = {
            "factual_consistency": "pass",
            "unsupported_claims": [],
            "missing_key_facts": [],
            "disposition_consistency": "pass",
            "disposition_inconsistencies": [],
            "criteria_alignment_issues": [],
            "missing_required_data_for_confident_interpretation": [
                "pH",
                "post-treatment reassessment",
            ],
            "mcg_admission_check": {
                "matched_criteria": ["diabetic_ketoacidosis"],
                "support_level": "high",
                "supported": True,
            },
            "needs_regeneration": False,
            "revision_instructions": [],
            "is_pass": True,
        }

        result = run_generate_pipeline("Glucose 320 mg/dL with dehydration.")

        self.assertCountEqual(
            result["structured_output"]["uncertainties_generated"],
            [
                "pH not available to fully assess severity of acidosis",
                "Serum osmolality not available",
                "Response to treatment not yet documented",
                "Post-treatment reassessment pending",
            ],
        )
        self.assertEqual(
            result["warning_groups"]["missing_data"],
            [
                "pH",
                "Serum osmolality",
                "Response to treatment documentation",
                "Post-treatment reassessment",
            ],
        )


class NoteParsingTests(SimpleTestCase):
    def test_parse_uploaded_docx_extracts_text(self):
        buffer = BytesIO()
        document = Document()
        document.add_paragraph("ER note text")
        document.add_paragraph("Second paragraph")
        document.save(buffer)
        buffer.seek(0)
        uploaded_file = SimpleUploadedFile(
            "er-note.docx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        result = parse_uploaded_note_file(uploaded_file)

        self.assertTrue(result["success"])
        self.assertIn("ER note text", result["text"])
        self.assertIsNone(result["error"])


class NoteParseUploadViewTests(TestCase):
    def setUp(self):
        self.client = APIClient()

    def test_parse_note_endpoint_rejects_invalid_extension(self):
        uploaded_file = SimpleUploadedFile(
            "invalid.txt",
            b"not supported",
            content_type="text/plain",
        )

        response = self.client.post(
            "/api/uploads/parse-note/",
            {"file": uploaded_file, "note_type": "er"},
            format="multipart",
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.data["success"])
        self.assertEqual(
            response.data["error"],
            "Unsupported file type. Please upload a PDF or DOCX file.",
        )

    def test_parse_note_endpoint_returns_warning_for_scanned_like_pdf(self):
        buffer = BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        writer.write(buffer)
        buffer.seek(0)
        uploaded_file = SimpleUploadedFile(
            "hp-note.pdf",
            buffer.getvalue(),
            content_type="application/pdf",
        )

        response = self.client.post(
            "/api/uploads/parse-note/",
            {"file": uploaded_file, "note_type": "hp"},
            format="multipart",
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.data["success"])
        self.assertEqual(response.data["text"], "")
        self.assertEqual(
            response.data["warning"],
            "No extractable text found in this PDF. It may be scanned.",
        )


class CaseDeleteViewTests(TestCase):
    def setUp(self):
        self.client = APIClient()

    def test_delete_case_returns_success_and_removes_case(self):
        from apps.cases.models import Case

        case = Case.objects.create(title="Delete me", original_note="note")

        response = self.client.delete(f"/api/cases/{case.id}/")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data, {"success": True})
        self.assertFalse(Case.objects.filter(id=case.id).exists())


class CaseTraceabilityPersistenceTests(TestCase):
    def setUp(self):
        self.client = APIClient()

    @patch("apps.cases.views.run_generate_pipeline")
    def test_generate_persists_traceability_metadata_and_retrieve_returns_it(
        self,
        mock_run_generate_pipeline,
    ):
        from apps.cases.models import Case

        case = Case.objects.create(title="Traceability", original_note="note text")
        mock_run_generate_pipeline.return_value = {
            "structured_output": {
                "chief_complaint_generated": "Hyperglycemia",
                "hpi_summary_generated": "Hyperglycemia requiring reassessment.",
                "key_findings_generated": ["Glucose 320 mg/dL"],
                "suspected_conditions_generated": ["Hyperglycemia"],
                "disposition_generated": "Admit",
                "uncertainties_generated": ["pH not available to fully assess severity of acidosis"],
            },
            "mcg_result": {
                "applicable": True,
                "support_level": "moderate",
                "supported": True,
            },
            "revised_hpi": "Generated revised HPI.",
            "verification": {
                "is_pass": True,
                "factual_consistency": "pass",
                "unsupported_claims": [],
                "missing_key_facts": [],
                "disposition_consistency": "pass",
                "disposition_inconsistencies": [],
                "criteria_alignment_issues": [],
                "missing_required_data_for_confident_interpretation": [],
                "mcg_admission_check": {
                    "applicable": True,
                    "matched_criteria": [],
                    "support_level": "moderate",
                    "supported": True,
                },
                "needs_regeneration": False,
                "revision_instructions": [],
            },
            "warnings": ["Missing data: pH"],
            "warning_groups": {
                "missing_data": ["pH"],
                "potential_issues": [],
            },
            "confidence_result": {
                "score": 0.74,
                "level": "Medium",
                "label": "Admission Support Confidence",
                "factors": [
                    {
                        "type": "positive",
                        "label": "Verifier passed",
                        "impact": 0.10,
                    }
                ],
            },
        }

        generate_response = self.client.post(f"/api/cases/{case.id}/generate/")

        self.assertEqual(generate_response.status_code, 200)
        self.assertEqual(
            generate_response.data["generated_result"]["generation_warnings"],
            ["Missing data: pH"],
        )
        self.assertEqual(
            generate_response.data["generated_result"]["mcg_result"]["support_level"],
            "moderate",
        )
        self.assertEqual(
            generate_response.data["generated_result"]["verification_result"]["factual_consistency"],
            "pass",
        )
        self.assertEqual(
            generate_response.data["generated_result"]["confidence_result"]["score"],
            0.74,
        )

        retrieve_response = self.client.get(f"/api/cases/{case.id}/")

        self.assertEqual(retrieve_response.status_code, 200)
        self.assertEqual(
            retrieve_response.data["generated_result"]["generation_warnings"],
            ["Missing data: pH"],
        )
        self.assertEqual(
            retrieve_response.data["generated_result"]["mcg_result"]["support_level"],
            "moderate",
        )
        self.assertEqual(
            retrieve_response.data["generated_result"]["verification_result"]["factual_consistency"],
            "pass",
        )
        self.assertEqual(
            retrieve_response.data["generated_result"]["confidence_result"]["score"],
            0.74,
        )
